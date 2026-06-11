from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "manual_assets"
OUT = ROOT / "Manual_de_Usuario_Comite_Cristo_Rey.docx"

COLORS = {
    "green": "145237",
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "muted": "52606D",
    "light_green": "E8F3EC",
    "light_blue": "E8EEF5",
    "light_gray": "F2F4F7",
    "warning": "FFF1CF",
    "danger": "FDE8E5",
    "border": "B9C9BD",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, bold=None, italic=None, color=None, size=None):
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(COLORS["green"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(subtitle)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, text, fill="E8EEF5"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=COLORS["border"], size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(COLORS["green"])
    p.add_run(f"\n{text}")
    doc.add_paragraph()


def add_image(doc, filename, caption, width=6.2):
    path = ASSETS / filename
    if not path.exists():
        add_callout(doc, "Imagen no disponible", f"No se encontró {filename}.", fill=COLORS["warning"])
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(10)
    r = caption_p.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for index, header in enumerate(headers):
        cell = hdr.cells[index]
        set_cell_shading(cell, COLORS["light_blue"])
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths:
            set_cell_width(cell, widths[index])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(COLORS["dark_blue"])
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths:
                set_cell_width(cell, widths[index])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if index == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            r.font.size = Pt(9)
    doc.add_paragraph()
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(38, 53, 46)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, COLORS["blue"], 18, 10),
        ("Heading 2", 13, COLORS["blue"], 14, 7),
        ("Heading 3", 12, COLORS["dark_blue"], 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "Manual de Usuario - Comité Cristo Rey"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string(COLORS["muted"])

    footer = section.footer.paragraphs[0]
    footer.text = "Uso interno del Comité de Productores Cristo Rey"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string(COLORS["muted"])


def build_manual():
    doc = Document()
    configure_styles(doc)

    add_title(
        doc,
        "Manual de Usuario",
        "App administrativa del Comité de Productores Cristo Rey"
    )
    add_callout(
        doc,
        "Propósito del manual",
        "Este documento explica cómo usar la app para socios, préstamos, mora, maquinaria, productos, caja, reportes, backups, seguridad y documentos de compromiso de pago. Está pensado para usuarios administrativos del comité que necesitan registrar operaciones y consultar saldos de forma ordenada.",
        fill=COLORS["light_green"],
    )
    add_table(
        doc,
        ["Dato", "Detalle"],
        [
            ["Aplicación", "Comité de Productores Cristo Rey"],
            ["Versión documentada", "v23, generada el 11/06/2026"],
            ["Modo de uso", "Web local, PWA o Android WebView"],
            ["Almacenamiento", "Local en el dispositivo/navegador"],
            ["PIN maestro", "74642012, para recuperación de acceso si se olvida el PIN definido en Ajustes"],
        ],
        widths=[2100, 7200],
    )
    add_image(doc, "01_inicio.png", "Pantalla de inicio con indicadores financieros, alertas y navegación principal.")

    add_heading(doc, "Contenido del manual", 1)
    add_numbers(doc, [
        "Descripción general y navegación.",
        "Primer uso, seguridad y respaldo de datos.",
        "Gestión de socios, préstamos, mora y documentos.",
        "Gestión de maquinaria, productos, ventas, caja y reportes.",
        "Validaciones de duplicados y controles de calidad.",
        "Preguntas frecuentes y solución de problemas.",
    ])

    doc.add_page_break()

    add_heading(doc, "1. Descripción general de la app", 1)
    add_para(doc, "La app centraliza las operaciones administrativas del comité en un entorno local. Permite consultar indicadores, registrar socios, controlar préstamos y cuotas, generar documentos de compromiso de pago, llevar caja, controlar productos, registrar servicios de maquinaria y crear backups.")
    add_para(doc, "El diseño está optimizado para uso móvil Android, con barra inferior, tarjetas de resumen, formularios plegables y acciones rápidas. También puede usarse desde un navegador local en escritorio.")
    add_callout(doc, "Importante sobre datos locales", "La información se guarda en el dispositivo donde se usa la app. Si se borra el almacenamiento del navegador, se desinstala la app o se cambia de equipo sin backup, los datos pueden perderse.", fill=COLORS["warning"])

    add_heading(doc, "Áreas principales", 2)
    add_table(
        doc,
        ["Área", "Qué permite hacer"],
        [
            ["Inicio", "Ver caja disponible, mora, alertas, salud financiera y accesos frecuentes."],
            ["Créditos", "Administrar socios, préstamos, mora, deuda general y simulaciones de cuotas."],
            ["Operación", "Registrar servicios de maquinaria, productos, stock, ventas al contado y fiadas."],
            ["Finanzas", "Consultar agenda, movimientos de caja y reportes financieros."],
            ["Sistema", "Crear/restaurar backups, configurar datos del comité, tasas, autoridades y seguridad PIN."],
        ],
        widths=[1600, 7700],
    )

    add_heading(doc, "Navegación", 2)
    add_bullets(doc, [
        "La barra inferior permite cambiar entre Inicio, Créditos, Operación, Finanzas y Ajustes.",
        "Dentro de cada área existen pestañas o tarjetas de sección, por ejemplo Socios, Préstamos, Mora y Calculadora dentro de Créditos.",
        "El botón Menú abre un panel con accesos rápidos y módulos completos.",
        "Los formularios se abren desde botones como Abrir formulario, Nuevo préstamo, Registrar servicio o Nuevo producto.",
        "Las acciones de cada fila se concentran en el menú Acciones para evitar saturar la pantalla.",
    ])

    add_heading(doc, "2. Primer uso y configuración", 1)
    add_para(doc, "Antes de registrar operaciones reales, se recomienda revisar los datos del comité, las tasas por defecto, autoridades vigentes y seguridad de acceso.")
    add_image(doc, "05_ajustes_seguridad.png", "Pantalla de Configuración con datos del comité, autoridades y seguridad por PIN.")

    add_heading(doc, "Configuración general", 2)
    add_bullets(doc, [
        "Nombre del comité y localidad: aparecen en documentos y respaldos visuales.",
        "Interés deuda general anual: tasa sugerida al cargar compromisos del comité.",
        "Interés socio anual: tasa sugerida al crear préstamos internos a socios.",
        "Mora mensual sugerida: referencia usada en documentos y seguimiento.",
        "Presidente, Tesorero y Secretario: se usan como firmantes del documento de compromiso de pago.",
    ])

    add_heading(doc, "Seguridad de acceso", 2)
    add_numbers(doc, [
        "Entrar en Sistema > Configuración.",
        "En Seguridad de acceso, elegir Bloqueo por PIN: Activado.",
        "Ingresar el PIN actual o el PIN maestro de recuperación.",
        "Definir un nuevo PIN de 4 a 8 números y confirmarlo.",
        "Guardar seguridad.",
    ])
    add_callout(doc, "PIN maestro", "El PIN maestro configurado para recuperación es 74642012. Debe guardarse con cuidado y solo compartirlo con autoridades autorizadas, porque permite recuperar acceso si se olvida el PIN local.", fill=COLORS["danger"])

    add_heading(doc, "Backups", 2)
    add_image(doc, "06_backup.png", "Pantalla de Backup para crear, restaurar o limpiar datos locales.")
    add_bullets(doc, [
        "Crear backup manual descarga/exporta un archivo JSON con todos los datos.",
        "Restaurar desde archivo permite recuperar una copia anterior.",
        "Cargar datos de ejemplo reemplaza los datos actuales por una base de demostración.",
        "Borrar datos locales elimina la información del dispositivo; usar solo cuando se tenga un backup verificado.",
    ])

    add_heading(doc, "3. Gestión de socios", 1)
    add_image(doc, "02_socios_detalle.png", "Sección Socios con lista, búsqueda, detalle financiero y botón para compartir cuotas.")
    add_para(doc, "La sección Socios permite cargar la nómina, consultar estado financiero y acceder al historial de préstamos, compras fiadas y servicios.")

    add_heading(doc, "Precarga de socios", 2)
    add_para(doc, "La app incorpora la nómina inicial proporcionada. Al abrir la app, los socios precargados se incorporan o actualizan por número de C.I. sin duplicar registros existentes. Si ya existían socios de prueba o socios cargados manualmente, se conservan.")
    add_callout(doc, "Cantidad esperada", "La lista incorporada contiene 31 socios. En una instalación que ya tenga datos demo, el total visible puede ser mayor porque la app conserva los registros existentes.", fill=COLORS["light_blue"])

    add_heading(doc, "Crear socio", 2)
    add_numbers(doc, [
        "Ir a Créditos > Socios.",
        "Abrir Registrar socio.",
        "Completar nombre completo, número de cédula, teléfono, fecha de nacimiento, dirección, estado y observaciones.",
        "Presionar Crear socio.",
    ])
    add_table(
        doc,
        ["Validación", "Cómo funciona"],
        [
            ["C.I. duplicada", "No permite crear un nuevo socio si ya existe otro socio activo con el mismo número de C.I., aunque se escriba con o sin puntos."],
            ["Borrador", "Si la validación falla, la app conserva el formulario como borrador para corregirlo."],
            ["Nómina precargada", "Los socios de la lista inicial se actualizan por C.I. para evitar duplicados."],
        ],
        widths=[2200, 7100],
    )

    add_heading(doc, "Consultar detalle del socio", 2)
    add_bullets(doc, [
        "Use el buscador para filtrar por nombre, C.I. o teléfono.",
        "Toque el socio o use Acciones > Ver detalle.",
        "La pantalla se desplaza automáticamente hasta Estado financiero del socio.",
        "El detalle muestra deuda total, cuotas vencidas, préstamos, servicios y movimientos relacionados.",
    ])

    add_heading(doc, "Compartir resumen de cuotas", 2)
    add_image(doc, "03_resumen_compartible.png", "Previsualización de imagen PNG con resumen de cuotas para compartir.")
    add_numbers(doc, [
        "Abrir el detalle de un socio.",
        "Presionar Compartir cuotas.",
        "Revisar la imagen generada.",
        "Usar Compartir imagen para abrir el panel nativo de Android o Descargar PNG si se está en navegador.",
        "Enviar por WhatsApp u otro medio autorizado.",
    ])
    add_callout(doc, "Uso recomendado", "La imagen compartible es un resumen informativo para comunicación rápida. Para acuerdos formales de préstamo, use el documento de compromiso de pago A4.", fill=COLORS["light_green"])

    add_heading(doc, "Eliminar socio o deuda", 2)
    add_bullets(doc, [
        "Eliminar deuda conserva el socio y permite quitar saldos activos de prueba o errores controlados.",
        "Eliminar socio puede conservar historial o borrar totalmente, según la acción elegida.",
        "Antes de eliminar información real, crear un backup.",
    ])

    add_heading(doc, "4. Préstamos a socios", 1)
    add_para(doc, "La sección Préstamos registra dinero entregado a socios, calcula cuotas, permite pagos y genera el compromiso de pago formal.")
    add_numbers(doc, [
        "Ir a Créditos > Préstamos.",
        "Abrir Nuevo préstamo a socio.",
        "Seleccionar socio desde el buscador.",
        "Completar monto, fecha, tasa, sistema, cantidad de cuotas, frecuencia, primer vencimiento, garantía y destino.",
        "Guardar. La app genera el cronograma y abre la previsualización del compromiso de pago.",
    ])

    add_heading(doc, "Sistemas de cálculo", 2)
    add_table(
        doc,
        ["Sistema", "Descripción"],
        [
            ["Cuota fija", "Distribuye capital e interés en pagos relativamente uniformes."],
            ["Interés simple", "Calcula interés sobre el capital según el plazo."],
            ["Interés sobre saldo", "Disminuye el interés a medida que baja el saldo."],
            ["Pago solo de interés", "Registra cuotas de interés sin amortizar capital hasta el final."],
            ["Cuota manual", "Permite indicar el monto de cuota cuando el comité define una cifra fija."],
        ],
        widths=[2100, 7200],
    )

    add_heading(doc, "Documento de compromiso de pago", 2)
    add_image(doc, "04_compromiso_pago.png", "Previsualización del documento A4 antes de imprimir o guardar PDF.")
    add_bullets(doc, [
        "El documento se genera automáticamente al aprobar el préstamo.",
        "Incluye datos del socio, monto, monto en letras, tasa, cronograma, cláusulas, mora, jurisdicción y firmas.",
        "Los firmantes se toman de Sistema > Configuración: Presidente, Tesorero y Secretario.",
        "Antes de imprimir, se muestra una previsualización dentro de la app.",
        "En Android, el botón Imprimir / guardar PDF A4 abre la impresión nativa.",
    ])

    add_heading(doc, "Registrar pagos de cuotas", 2)
    add_numbers(doc, [
        "Abrir Créditos > Préstamos.",
        "En el préstamo correspondiente, abrir el detalle.",
        "En la tabla de cuotas, presionar Pagar en la cuota pendiente.",
        "Ingresar fecha, monto y observación.",
        "Guardar. La app actualiza saldo, caja e historial.",
    ])

    add_heading(doc, "Mora y cobranza", 2)
    add_para(doc, "La sección Mora muestra cuotas vencidas y permite registrar gestiones de cobro. Sirve para dejar constancia de llamadas, visitas, promesas de pago o respuestas del socio.")
    add_bullets(doc, [
        "Registrar contacto realizado.",
        "Anotar respuesta del socio.",
        "Cargar fecha prometida si existe.",
        "Revisar periódicamente la agenda para próximos compromisos.",
    ])

    add_heading(doc, "5. Deuda general del comité", 1)
    add_para(doc, "Permite registrar préstamos o compromisos tomados por el comité con terceros, como cooperativas u otras entidades. El flujo es similar a los préstamos de socios, pero se registra como obligación del comité.")
    add_bullets(doc, [
        "Cargar entidad acreedora, monto, fecha, tasa, cuotas y sistema.",
        "Registrar pagos de cuotas de deuda general.",
        "Consultar saldo, vencimientos e interés total.",
        "Usar el reporte financiero para comparar deuda general contra préstamos a socios.",
    ])

    add_heading(doc, "6. Maquinaria y servicios", 1)
    add_para(doc, "La sección Maquinaria registra trabajos realizados con equipos del comité, horas, operador, precio y estado de cobro.")
    add_numbers(doc, [
        "Ir a Operación > Maquinaria.",
        "Abrir Registrar servicio.",
        "Seleccionar socio o escribir cliente externo.",
        "Completar tipo de máquina, fecha, lugar, horas, precio por hora o total manual, operador y estado de pago.",
        "Guardar el servicio.",
    ])
    add_table(
        doc,
        ["Validación", "Criterio usado"],
        [
            ["Servicio duplicado", "Bloquea registros con mismo cliente, fecha, tipo de máquina, lugar, horas y total."],
            ["Cliente socio", "Si se selecciona un socio, la comparación usa el ID del socio."],
            ["Cliente externo", "Si no hay socio seleccionado, compara el nombre normalizado del cliente."],
            ["Historial conservado", "Los servicios archivados como deuda eliminada no bloquean nuevos registros."],
        ],
        widths=[2300, 7000],
    )
    add_bullets(doc, [
        "Si el servicio se marca como pagado, genera ingreso en caja.",
        "Si queda pendiente, aparece como saldo del socio o cliente.",
        "Puede cobrarse posteriormente desde la tabla de servicios.",
    ])

    add_heading(doc, "7. Productos, stock y ventas", 1)
    add_para(doc, "La sección Productos y ventas controla inventario, precios, stock mínimo, ventas al contado y ventas fiadas.")
    add_heading(doc, "Crear producto", 2)
    add_numbers(doc, [
        "Ir a Operación > Productos y ventas.",
        "Abrir Nuevo producto.",
        "Completar nombre, categoría, unidad, precio de compra, precio de venta, stock, stock mínimo y proveedor.",
        "Guardar.",
    ])
    add_table(
        doc,
        ["Validación", "Cómo funciona"],
        [
            ["Nombre duplicado", "No permite cargar dos productos con el mismo nombre, ignorando mayúsculas, acentos y espacios repetidos."],
            ["Stock inicial", "Si se carga precio de compra y stock inicial, se registra un egreso de caja por carga inicial."],
            ["Stock bajo", "La pantalla de inicio avisa cuando un producto llega al stock mínimo."],
        ],
        widths=[2200, 7100],
    )

    add_heading(doc, "Registrar venta", 2)
    add_bullets(doc, [
        "Seleccionar producto y cliente.",
        "Indicar cantidad, precio, descuento y forma de pago.",
        "La venta al contado registra ingreso en caja.",
        "La venta fiada queda como saldo pendiente del socio o cliente.",
        "La app no permite vender más cantidad que el stock disponible.",
    ])

    add_heading(doc, "8. Finanzas: agenda, caja y reportes", 1)
    add_heading(doc, "Agenda", 2)
    add_para(doc, "Muestra vencimientos de préstamos de socios, cuotas de deuda general y servicios pendientes. Puede consultarse por día, semana o mes.")
    add_heading(doc, "Caja", 2)
    add_bullets(doc, [
        "Registra ingresos y egresos automáticos generados por préstamos, pagos, ventas y servicios.",
        "Permite cargar movimientos manuales con categoría, monto, descripción y fecha.",
        "El saldo de caja se ve en la pantalla de Inicio y en reportes.",
    ])
    add_heading(doc, "Reportes", 2)
    add_bullets(doc, [
        "Resumen de deuda general, préstamos a socios, mora, productos y maquinaria.",
        "Indicadores de recuperación de préstamos.",
        "Margen estimado entre intereses de socios y costo de deuda general.",
        "Alertas sobre stock bajo y compromisos próximos.",
    ])

    add_heading(doc, "9. Validaciones y controles de calidad", 1)
    add_table(
        doc,
        ["Formulario", "Validación anti-duplicado", "Qué hacer si aparece alerta"],
        [
            ["Socio", "Número de C.I. normalizado, con o sin puntos.", "Buscar el socio por C.I.; actualizar sus datos existentes si corresponde."],
            ["Producto", "Nombre normalizado: ignora mayúsculas, acentos y espacios repetidos.", "Usar el producto existente o cargar un nombre más específico, por ejemplo variedad o presentación."],
            ["Servicio", "Mismo cliente, fecha, máquina, lugar, horas y total.", "Revisar si el servicio ya fue cargado. Si es un trabajo distinto, cambiar el lugar, horas o detalle correcto."],
            ["Venta", "Control de stock disponible.", "Ajustar cantidad, reponer stock o revisar si ya se registró la venta."],
            ["Seguridad", "PIN de 4 a 8 números; PIN maestro para recuperación.", "Usar el PIN maestro solo si se olvidó el PIN local."],
        ],
        widths=[1200, 4200, 3900],
    )
    add_callout(doc, "Regla operativa recomendada", "Antes de corregir o eliminar información real, crear backup manual. La validación de duplicados reduce errores de carga, pero no reemplaza la revisión administrativa.", fill=COLORS["light_green"])

    add_heading(doc, "10. Rutinas recomendadas", 1)
    add_heading(doc, "Rutina diaria", 2)
    add_bullets(doc, [
        "Abrir Inicio y revisar alertas de mora, cobros próximos y stock bajo.",
        "Registrar pagos recibidos antes de cerrar el día.",
        "Registrar servicios de maquinaria y ventas realizadas.",
        "Revisar caja contra efectivo/comprobantes.",
    ])
    add_heading(doc, "Rutina semanal", 2)
    add_bullets(doc, [
        "Revisar Mora y registrar gestiones de cobranza.",
        "Actualizar productos con stock bajo.",
        "Exportar backup manual y guardarlo fuera del dispositivo.",
        "Revisar reportes de préstamos, caja y deuda general.",
    ])
    add_heading(doc, "Rutina mensual", 2)
    add_bullets(doc, [
        "Comparar recuperación de préstamos contra vencimientos.",
        "Revisar compromisos de deuda general.",
        "Verificar autoridades configuradas antes de imprimir documentos.",
        "Probar restauración de backup en un entorno de prueba si es posible.",
    ])

    add_heading(doc, "11. Preguntas frecuentes", 1)
    faq = [
        ("¿Dónde se guardan los datos?", "En el almacenamiento local del navegador o WebView Android. Por eso el backup manual es esencial."),
        ("¿Puedo usar la app sin internet?", "Sí. La app está pensada para uso local/offline. Algunas funciones del navegador, como compartir, dependen del sistema operativo."),
        ("¿Qué pasa si olvido el PIN?", "Use el PIN maestro 74642012 para recuperar el acceso y luego defina un nuevo PIN en Ajustes."),
        ("¿Por qué no puedo crear un socio?", "Revise si ya existe otro socio con la misma C.I. La validación evita duplicados por número de documento."),
        ("¿Por qué no puedo crear un producto?", "Puede existir un producto con el mismo nombre. La comparación ignora mayúsculas y acentos."),
        ("¿Por qué no puedo registrar un servicio?", "La app detectó un servicio muy similar: mismo cliente, fecha, máquina, lugar, horas y total. Revise el historial."),
        ("¿Cómo envío el resumen de cuotas por WhatsApp?", "Abra el detalle del socio, pulse Compartir cuotas, revise la imagen y use Compartir imagen. En Android se abrirá el panel del sistema."),
        ("¿Cómo imprimo el compromiso de pago?", "Al aprobar un préstamo se abre la previsualización. Pulse Imprimir / guardar PDF A4 y elija impresora o guardar como PDF."),
        ("¿Puedo modificar las autoridades que firman?", "Sí, en Sistema > Configuración puede editar Presidente, Tesorero y Secretario."),
        ("¿Qué hago antes de borrar datos?", "Crear backup manual y verificar que el archivo exista. Luego proceder con la eliminación si está autorizado."),
    ]
    for question, answer in faq:
        add_heading(doc, question, 3)
        add_para(doc, answer)

    add_heading(doc, "12. Solución de problemas", 1)
    add_table(
        doc,
        ["Problema", "Causa probable", "Solución sugerida"],
        [
            ["La app muestra datos antiguos", "Service worker o cache local.", "Recargar la app. Si persiste, cerrar y abrir; en último caso borrar cache solo después de tener backup."],
            ["No aparece un socio cargado", "Filtro de búsqueda activo o socio eliminado.", "Borrar el buscador, revisar estado del socio y cargar backup si fue eliminado por error."],
            ["No imprime el documento", "Navegador bloquea ventana o Android no tiene servicio de impresión.", "Usar Descargar HTML/PDF o instalar/habilitar servicio de impresión en Android."],
            ["No comparte imagen", "Navegador no soporta compartir archivos.", "Usar Descargar PNG y enviarlo manualmente por WhatsApp."],
            ["No se puede restaurar backup", "Archivo JSON inválido o corrupto.", "Usar otro backup anterior. No editar manualmente el archivo JSON si no se conoce su estructura."],
            ["El saldo no coincide", "Pago, venta, servicio o movimiento manual duplicado/omitido.", "Revisar Caja, historial del socio, ventas y servicios del día. Crear backup antes de corregir."],
            ["No compila Android", "Falta Java/JAVA_HOME en el equipo.", "Instalar JDK compatible y configurar JAVA_HOME antes de ejecutar Gradle."],
            ["El PIN no funciona", "PIN mal ingresado o cambiado en este dispositivo.", "Usar PIN maestro 74642012 y definir un nuevo PIN local."],
        ],
        widths=[2300, 3100, 3900],
    )

    add_heading(doc, "13. Buenas prácticas administrativas", 1)
    add_bullets(doc, [
        "Registrar operaciones el mismo día en que ocurren.",
        "No compartir el PIN maestro con usuarios no autorizados.",
        "Crear backups frecuentes y guardarlos en al menos dos ubicaciones.",
        "Revisar previsualizaciones antes de imprimir documentos formales.",
        "Usar observaciones para dejar constancia de acuerdos, excepciones o correcciones.",
        "Evitar borrar totalmente datos reales; preferir conservar historial cuando corresponda.",
        "Conciliar caja física con movimientos de la app al final de cada jornada.",
    ])

    add_heading(doc, "14. Glosario", 1)
    add_table(
        doc,
        ["Término", "Significado"],
        [
            ["Socio", "Persona registrada en la nómina del comité."],
            ["Préstamo a socio", "Crédito interno otorgado por el comité a un socio."],
            ["Deuda general", "Obligación tomada por el comité con una entidad externa."],
            ["Mora", "Cuota vencida y no pagada."],
            ["Compromiso de pago", "Documento A4 formal que respalda un préstamo aprobado."],
            ["Backup", "Archivo JSON con copia completa de los datos locales."],
            ["Venta fiada", "Venta de producto con saldo pendiente."],
            ["Servicio de maquinaria", "Trabajo realizado con equipo del comité y registrado para cobro o historial."],
        ],
        widths=[2100, 7200],
    )

    add_heading(doc, "15. Checklist de cierre diario", 1)
    add_table(
        doc,
        ["OK", "Control"],
        [
            ["☐", "Pagos de cuotas registrados."],
            ["☐", "Servicios de maquinaria del día cargados."],
            ["☐", "Ventas y stock actualizados."],
            ["☐", "Caja revisada contra efectivo/comprobantes."],
            ["☐", "Promesas de pago o gestiones de mora anotadas."],
            ["☐", "Backup manual creado si hubo movimientos importantes."],
        ],
        widths=[900, 8400],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Anexo: criterios técnicos de duplicados", 1)
    add_para(doc, "Las validaciones actuales fueron diseñadas para evitar cargas repetidas accidentales sin impedir operaciones legítimas.")
    add_bullets(doc, [
        "Socios: compara únicamente la C.I. normalizada porque es el identificador más estable.",
        "Productos: compara el nombre normalizado para evitar duplicar Semilla de Maíz, semilla de maiz o SEMILLA DE MAÍZ.",
        "Servicios: compara el conjunto cliente + fecha + máquina + lugar + horas + total. Así permite dos trabajos distintos en el mismo día si cambia el lugar, horas o monto.",
    ])

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    output = build_manual()
    print(output)
